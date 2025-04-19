import os
import re
import tkinter as tk
from tkinter import filedialog, ttk, messagebox, font
import shutil
from docx import Document
import win32com.client
import threading
import pythoncom
import os.path
import time
import sys
import pandas as pd
from openpyxl import load_workbook

class SearchReplaceApp:
    def __init__(self, root):
        self.root = root
        self.root.title("Tìm kiếm và Thay thế trong tài liệu Word")
        # Đặt kích thước lớn hơn cho ứng dụng và cho phép thay đổi kích thước
        self.root.geometry("1000x750")
        self.root.minsize(900, 700)
        self.root.resizable(True, True)
        
        # Tạo phông chữ lớn hơn cho tiêu đề
        self.title_font = font.Font(size=12, weight="bold")
        
        # Biến lưu trữ đường dẫn files và thư mục
        self.selected_files = []
        self.selected_directory = ""
        self.search_text = ""
        self.replace_text = ""
        self.failed_files = []
        self.success_count = 0
        self.skip_count = 0
        
        # Dữ liệu cho tìm kiếm và thay thế từ bảng
        self.replacement_pairs = []
        self.rules_file_path = ""
        self.filename_replacement_pairs = []
        
        # Tạo giao diện
        self.create_widgets()
        
        # Kiểm tra kích thước màn hình và hiển thị cảnh báo nếu cần
        self.root.update_idletasks()
        self.check_screen_size()
        
        # Thiết lập sự kiện thay đổi kích thước để kiểm tra kích thước màn hình
        self.root.bind("<Configure>", lambda e: self.check_screen_size())
    
    def check_screen_size(self):
        """Kiểm tra kích thước cửa sổ và hiển thị cảnh báo nếu quá nhỏ"""
        width = self.root.winfo_width()
        height = self.root.winfo_height()
        
        if width < 900 or height < 700:
            self.screen_warning_label.config(text="Cảnh báo: Kích thước cửa sổ quá nhỏ, một số phần có thể không hiển thị đầy đủ.")
            self.screen_warning_label.grid(row=0, column=0, sticky="ew")
        else:
            self.screen_warning_label.grid_forget()

    # Phần còn lại của class được khai báo ở các file khác
    def create_widgets(self):
        # Frame chính
        main_frame = ttk.Frame(self.root, padding="10")
        main_frame.pack(fill=tk.BOTH, expand=True)
        
        # Label cảnh báo kích thước màn hình
        warning_frame = ttk.Frame(main_frame)
        warning_frame.pack(fill=tk.X)
        self.screen_warning_label = ttk.Label(warning_frame, 
                                             text="", 
                                             foreground="red", 
                                             justify=tk.CENTER)
        
        # Notebook (tabs)
        self.notebook = ttk.Notebook(main_frame)
        self.notebook.pack(fill=tk.BOTH, expand=True)
        
        # Tab 1: Tìm kiếm và thay thế đơn giản
        simple_tab = ttk.Frame(self.notebook)
        self.notebook.add(simple_tab, text="Thay thế trong nội dung")
        
        # Tab 2: Thay thế từ bảng
        table_tab = ttk.Frame(self.notebook)
        self.notebook.add(table_tab, text="Thay thế từ bảng quy tắc")
        
        # Tab 3: Thay thế trong tên file
        filename_tab = ttk.Frame(self.notebook)
        self.notebook.add(filename_tab, text="Thay thế trong tên file")
        
        # Tab 4: Thay thế tên file từ bảng
        filename_table_tab = ttk.Frame(self.notebook)
        self.notebook.add(filename_table_tab, text="Thay thế tên file từ bảng")
        
        # Thiết lập tab 1: Tìm kiếm và thay thế nội dung đơn giản
        self.setup_simple_tab(simple_tab)
        
        # Thiết lập tab 2: Thay thế nội dung từ bảng
        self.setup_table_tab(table_tab)
        
        # Thiết lập tab 3: Thay thế tên file đơn giản
        self.setup_filename_tab(filename_tab)
        
        # Thiết lập tab 4: Thay thế tên file từ bảng
        self.setup_filename_table_tab(filename_table_tab)
        
        # Khu vực hiển thị log (dùng chung cho cả bốn tab)
        log_frame = ttk.LabelFrame(main_frame, text="Log xử lý", padding="10")
        log_frame.pack(fill=tk.BOTH, expand=True, pady=5)
        
        # Thanh cuộn cho log
        log_scroll = ttk.Scrollbar(log_frame)
        log_scroll.pack(side=tk.RIGHT, fill=tk.Y)
        
        # Text widget cho log
        self.log_text = tk.Text(log_frame, wrap=tk.WORD, height=12, yscrollcommand=log_scroll.set)
        self.log_text.pack(fill=tk.BOTH, expand=True)
        log_scroll.config(command=self.log_text.yview)
    def setup_simple_tab(self, parent):
        # Frame cho phần chọn files/thư mục
        select_frame = ttk.LabelFrame(parent, text="Chọn Files hoặc Thư mục", padding="10")
        select_frame.pack(fill=tk.X, pady=5)
        
        # Nút chọn files
        select_files_button = ttk.Button(select_frame, text="Chọn Files", command=self.select_files)
        select_files_button.grid(row=0, column=0, padx=5, pady=5, sticky="w")
        
        # Nút chọn thư mục
        select_dir_button = ttk.Button(select_frame, text="Chọn Thư mục", command=self.select_directory)
        select_dir_button.grid(row=0, column=1, padx=5, pady=5, sticky="w")
        
        # Hiển thị đường dẫn đã chọn
        self.path_var = tk.StringVar()
        path_label = ttk.Label(select_frame, textvariable=self.path_var, wraplength=800)
        path_label.grid(row=1, column=0, columnspan=2, padx=5, pady=5, sticky="w")
        
        # Frame cho tìm kiếm và thay thế
        search_frame = ttk.LabelFrame(parent, text="Tìm kiếm và Thay thế", padding="10")
        search_frame.pack(fill=tk.X, pady=5)
        
        # Nhập chuỗi cần tìm
        ttk.Label(search_frame, text="Nhập chuỗi cần tìm:").grid(row=0, column=0, padx=5, pady=5, sticky="w")
        self.search_entry = ttk.Entry(search_frame, width=70)
        self.search_entry.grid(row=0, column=1, padx=5, pady=5, sticky="w")
        
        # Nhập chuỗi thay thế
        ttk.Label(search_frame, text="Nhập chuỗi thay thế:").grid(row=1, column=0, padx=5, pady=5, sticky="w")
        self.replace_entry = ttk.Entry(search_frame, width=70)
        self.replace_entry.grid(row=1, column=1, padx=5, pady=5, sticky="w")
        
        # Lựa chọn loại file
        ttk.Label(search_frame, text="Xử lý file:").grid(row=2, column=0, padx=5, pady=5, sticky="w")
        self.file_type_frame = ttk.Frame(search_frame)
        self.file_type_frame.grid(row=2, column=1, padx=5, pady=5, sticky="w")
        
        self.process_docx_var = tk.BooleanVar(value=True)
        self.process_doc_var = tk.BooleanVar(value=True)
        
        ttk.Checkbutton(self.file_type_frame, text=".docx", variable=self.process_docx_var).pack(side=tk.LEFT, padx=5)
        ttk.Checkbutton(self.file_type_frame, text=".doc", variable=self.process_doc_var).pack(side=tk.LEFT, padx=5)
        
        # Nút thực hiện
        execute_frame = ttk.Frame(parent)
        execute_frame.pack(fill=tk.X, pady=10)
        
        self.execute_button = ttk.Button(
            execute_frame, 
            text="Thực hiện tìm kiếm và thay thế", 
            command=self.execute_search_replace
        )
        self.execute_button.pack(side=tk.RIGHT, padx=5)
        
        # Thanh tiến trình
        progress_frame = ttk.Frame(parent)
        progress_frame.pack(fill=tk.X, pady=5)
        
        ttk.Label(progress_frame, text="Tiến trình:").pack(side=tk.LEFT, padx=5)
        self.progress = ttk.Progressbar(progress_frame, orient=tk.HORIZONTAL, length=800, mode='determinate')
        self.progress.pack(side=tk.LEFT, padx=5, fill=tk.X, expand=True)
        
        # Trạng thái
        status_frame = ttk.Frame(parent)
        status_frame.pack(fill=tk.X, pady=5)
        
        self.status_var = tk.StringVar(value="Sẵn sàng")
        status_label = ttk.Label(status_frame, textvariable=self.status_var, wraplength=800)
        status_label.pack(side=tk.LEFT, padx=5, fill=tk.X, expand=True)
    
    def setup_table_tab(self, parent):
        # Frame cho phần chọn file quy tắc thay thế
        rules_frame = ttk.LabelFrame(parent, text="Bước 1: Chọn tệp chứa bảng quy tắc thay thế", padding="10")
        rules_frame.pack(fill=tk.X, pady=5)
        
        # Nút chọn file quy tắc
        select_rules_button = ttk.Button(rules_frame, text="Chọn tệp Word/Excel", command=self.select_rules_file)
        select_rules_button.grid(row=0, column=0, padx=5, pady=5, sticky="w")
        
        # Hiển thị đường dẫn file quy tắc đã chọn
        self.rules_path_var = tk.StringVar()
        rules_path_label = ttk.Label(rules_frame, textvariable=self.rules_path_var, wraplength=800)
        rules_path_label.grid(row=1, column=0, columnspan=2, padx=5, pady=5, sticky="w")
        
        # Frame cho hiển thị quy tắc thay thế
        rules_display_frame = ttk.LabelFrame(parent, text="Danh sách quy tắc thay thế", padding="10")
        rules_display_frame.pack(fill=tk.BOTH, expand=True, pady=5)
        
        # Tạo Treeview để hiển thị danh sách quy tắc
        columns = ('stt', 'search', 'replace')
        self.rules_tree = ttk.Treeview(rules_display_frame, columns=columns, show='headings', height=8)
        
        # Định nghĩa headings
        self.rules_tree.heading('stt', text='STT')
        self.rules_tree.heading('search', text='Tìm kiếm')
        self.rules_tree.heading('replace', text='Thay thế')
        
        # Định nghĩa columns
        self.rules_tree.column('stt', width=50)
        self.rules_tree.column('search', width=400)
        self.rules_tree.column('replace', width=400)
        
        # Thêm thanh cuộn
        rules_scroll = ttk.Scrollbar(rules_display_frame, orient=tk.VERTICAL, command=self.rules_tree.yview)
        self.rules_tree.configure(yscrollcommand=rules_scroll.set)
        
        # Pack các widgets
        rules_scroll.pack(side=tk.RIGHT, fill=tk.Y)
        self.rules_tree.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        
        # Frame cho phần chọn files/thư mục cần xử lý
        target_frame = ttk.LabelFrame(parent, text="Bước 2: Chọn Files hoặc Thư mục cần xử lý", padding="10")
        target_frame.pack(fill=tk.X, pady=5)
        
        # Nút chọn files
        select_target_files_button = ttk.Button(target_frame, text="Chọn Files", command=self.select_target_files)
        select_target_files_button.grid(row=0, column=0, padx=5, pady=5, sticky="w")
        
        # Nút chọn thư mục
        select_target_dir_button = ttk.Button(target_frame, text="Chọn Thư mục", command=self.select_target_directory)
        select_target_dir_button.grid(row=0, column=1, padx=5, pady=5, sticky="w")
        
        # Hiển thị đường dẫn đã chọn
        self.target_path_var = tk.StringVar()
        target_path_label = ttk.Label(target_frame, textvariable=self.target_path_var, wraplength=800)
        target_path_label.grid(row=1, column=0, columnspan=2, padx=5, pady=5, sticky="w")
        
        # Lựa chọn loại file
        ttk.Label(target_frame, text="Xử lý file:").grid(row=2, column=0, padx=5, pady=5, sticky="w")
        self.target_file_type_frame = ttk.Frame(target_frame)
        self.target_file_type_frame.grid(row=2, column=1, padx=5, pady=5, sticky="w")
        
        self.target_process_docx_var = tk.BooleanVar(value=True)
        self.target_process_doc_var = tk.BooleanVar(value=True)
        
        ttk.Checkbutton(self.target_file_type_frame, text=".docx", variable=self.target_process_docx_var).pack(side=tk.LEFT, padx=5)
        ttk.Checkbutton(self.target_file_type_frame, text=".doc", variable=self.target_process_doc_var).pack(side=tk.LEFT, padx=5)
        
        # Frame cho thực hiện
        target_execute_frame = ttk.LabelFrame(parent, text="Bước 3: Thực hiện", padding="10")
        target_execute_frame.pack(fill=tk.X, pady=10)
        
        # Nút thực hiện với tên nhất quán như tab 1
        self.target_execute_button = ttk.Button(
            target_execute_frame, 
            text="Thực hiện tìm kiếm và thay thế", 
            command=self.execute_table_search_replace
        )
        self.target_execute_button.pack(side=tk.RIGHT, padx=5, pady=5)
        
        # Thanh tiến trình
        target_progress_frame = ttk.Frame(target_execute_frame)
        target_progress_frame.pack(fill=tk.X, pady=5)
        
        ttk.Label(target_progress_frame, text="Tiến trình:").pack(side=tk.LEFT, padx=5)
        self.target_progress = ttk.Progressbar(target_progress_frame, orient=tk.HORIZONTAL, length=800, mode='determinate')
        self.target_progress.pack(side=tk.LEFT, padx=5, fill=tk.X, expand=True)
        
        # Trạng thái
        target_status_frame = ttk.Frame(target_execute_frame)
        target_status_frame.pack(fill=tk.X, pady=5)
        
        self.target_status_var = tk.StringVar(value="Sẵn sàng")
        target_status_label = ttk.Label(target_status_frame, textvariable=self.target_status_var, wraplength=800)
        target_status_label.pack(side=tk.LEFT, padx=5, fill=tk.X, expand=True)
    def setup_filename_tab(self, parent):
        # Frame cho phần chọn files/thư mục
        select_frame = ttk.LabelFrame(parent, text="Chọn Files hoặc Thư mục", padding="10")
        select_frame.pack(fill=tk.X, pady=5)
        
        # Nút chọn files
        select_files_button = ttk.Button(select_frame, text="Chọn Files", command=self.select_files_for_rename)
        select_files_button.grid(row=0, column=0, padx=5, pady=5, sticky="w")
        
        # Nút chọn thư mục
        select_dir_button = ttk.Button(select_frame, text="Chọn Thư mục", command=self.select_directory_for_rename)
        select_dir_button.grid(row=0, column=1, padx=5, pady=5, sticky="w")
        
        # Hiển thị đường dẫn đã chọn
        self.filename_path_var = tk.StringVar()
        path_label = ttk.Label(select_frame, textvariable=self.filename_path_var, wraplength=800)
        path_label.grid(row=1, column=0, columnspan=2, padx=5, pady=5, sticky="w")
        
        # Tùy chọn include subdirectories
        self.filename_include_subdirs_var = tk.BooleanVar(value=True)
        include_subdirs_check = ttk.Checkbutton(
            select_frame, 
            text="Bao gồm các thư mục con",
            variable=self.filename_include_subdirs_var
        )
        include_subdirs_check.grid(row=2, column=0, columnspan=2, padx=5, pady=5, sticky="w")
        
        # Frame cho tìm kiếm và thay thế
        search_frame = ttk.LabelFrame(parent, text="Tìm kiếm và Thay thế trong tên file", padding="10")
        search_frame.pack(fill=tk.X, pady=5)
        
        # Nhập chuỗi cần tìm
        ttk.Label(search_frame, text="Nhập chuỗi cần tìm:").grid(row=0, column=0, padx=5, pady=5, sticky="w")
        self.filename_search_entry = ttk.Entry(search_frame, width=70)
        self.filename_search_entry.grid(row=0, column=1, padx=5, pady=5, sticky="w")
        
        # Nhập chuỗi thay thế
        ttk.Label(search_frame, text="Nhập chuỗi thay thế:").grid(row=1, column=0, padx=5, pady=5, sticky="w")
        self.filename_replace_entry = ttk.Entry(search_frame, width=70)
        self.filename_replace_entry.grid(row=1, column=1, padx=5, pady=5, sticky="w")
        
        # Lựa chọn loại file
        ttk.Label(search_frame, text="Xử lý file:").grid(row=2, column=0, padx=5, pady=5, sticky="w")
        self.filename_type_frame = ttk.Frame(search_frame)
        self.filename_type_frame.grid(row=2, column=1, padx=5, pady=5, sticky="w")
        
        self.filename_process_docx_var = tk.BooleanVar(value=True)
        self.filename_process_doc_var = tk.BooleanVar(value=True)
        self.filename_process_all_var = tk.BooleanVar(value=False)
        
        ttk.Checkbutton(self.filename_type_frame, text=".docx", variable=self.filename_process_docx_var).pack(side=tk.LEFT, padx=5)
        ttk.Checkbutton(self.filename_type_frame, text=".doc", variable=self.filename_process_doc_var).pack(side=tk.LEFT, padx=5)
        ttk.Checkbutton(self.filename_type_frame, text="Tất cả các file", variable=self.filename_process_all_var, 
                        command=self.toggle_all_files).pack(side=tk.LEFT, padx=5)
        
        # Tùy chọn xem trước thay đổi
        preview_frame = ttk.LabelFrame(parent, text="Xem trước kết quả", padding="10")
        preview_frame.pack(fill=tk.BOTH, expand=True, pady=5)
        
        # Tạo Treeview để hiển thị danh sách tên file
        columns = ('old_name', 'new_name', 'path')
        self.filename_preview_tree = ttk.Treeview(preview_frame, columns=columns, show='headings', height=8)
        
        # Định nghĩa headings
        self.filename_preview_tree.heading('old_name', text='Tên cũ')
        self.filename_preview_tree.heading('new_name', text='Tên mới')
        self.filename_preview_tree.heading('path', text='Đường dẫn')
        
        # Định nghĩa columns
        self.filename_preview_tree.column('old_name', width=300)
        self.filename_preview_tree.column('new_name', width=300)
        self.filename_preview_tree.column('path', width=300)
        
        # Thêm thanh cuộn
        preview_scroll = ttk.Scrollbar(preview_frame, orient=tk.VERTICAL, command=self.filename_preview_tree.yview)
        self.filename_preview_tree.configure(yscrollcommand=preview_scroll.set)
        
        # Pack các widgets
        preview_scroll.pack(side=tk.RIGHT, fill=tk.Y)
        self.filename_preview_tree.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        
        # Nút xem trước và thực hiện
        button_frame = ttk.Frame(parent)
        button_frame.pack(fill=tk.X, pady=10)
        
        self.filename_preview_button = ttk.Button(
            button_frame, 
            text="Xem trước thay đổi", 
            command=self.preview_filename_changes
        )
        self.filename_preview_button.pack(side=tk.LEFT, padx=5)
        
        self.filename_execute_button = ttk.Button(
            button_frame, 
            text="Thực hiện đổi tên", 
            command=self.execute_filename_replace
        )
        self.filename_execute_button.pack(side=tk.RIGHT, padx=5)
        
        # Thanh tiến trình
        progress_frame = ttk.Frame(parent)
        progress_frame.pack(fill=tk.X, pady=5)
        
        ttk.Label(progress_frame, text="Tiến trình:").pack(side=tk.LEFT, padx=5)
        self.filename_progress = ttk.Progressbar(progress_frame, orient=tk.HORIZONTAL, length=800, mode='determinate')
        self.filename_progress.pack(side=tk.LEFT, padx=5, fill=tk.X, expand=True)
        
        # Trạng thái
        status_frame = ttk.Frame(parent)
        status_frame.pack(fill=tk.X, pady=5)
        
        self.filename_status_var = tk.StringVar(value="Sẵn sàng")
        status_label = ttk.Label(status_frame, textvariable=self.filename_status_var, wraplength=800)
        status_label.pack(side=tk.LEFT, padx=5, fill=tk.X, expand=True)
    
    def setup_filename_table_tab(self, parent):
        # Frame cho phần chọn file quy tắc thay thế
        rules_frame = ttk.LabelFrame(parent, text="Bước 1: Chọn tệp chứa bảng quy tắc thay thế tên file", padding="10")
        rules_frame.pack(fill=tk.X, pady=5)
        
        # Nút chọn file quy tắc
        select_rules_button = ttk.Button(rules_frame, text="Chọn tệp Word/Excel", command=self.select_filename_rules_file)
        select_rules_button.grid(row=0, column=0, padx=5, pady=5, sticky="w")
        
        # Hiển thị đường dẫn file quy tắc đã chọn
        self.filename_rules_path_var = tk.StringVar()
        rules_path_label = ttk.Label(rules_frame, textvariable=self.filename_rules_path_var, wraplength=800)
        rules_path_label.grid(row=1, column=0, columnspan=2, padx=5, pady=5, sticky="w")
        
        # Frame cho hiển thị quy tắc thay thế
        rules_display_frame = ttk.LabelFrame(parent, text="Danh sách quy tắc thay thế tên file", padding="10")
        rules_display_frame.pack(fill=tk.BOTH, expand=True, pady=5)
        
        # Tạo Treeview để hiển thị danh sách quy tắc
        columns = ('stt', 'search', 'replace')
        self.filename_rules_tree = ttk.Treeview(rules_display_frame, columns=columns, show='headings', height=6)
        
        # Định nghĩa headings
        self.filename_rules_tree.heading('stt', text='STT')
        self.filename_rules_tree.heading('search', text='Tìm kiếm')
        self.filename_rules_tree.heading('replace', text='Thay thế')
        
        # Định nghĩa columns
        self.filename_rules_tree.column('stt', width=50)
        self.filename_rules_tree.column('search', width=400)
        self.filename_rules_tree.column('replace', width=400)
        
        # Thêm thanh cuộn
        rules_scroll = ttk.Scrollbar(rules_display_frame, orient=tk.VERTICAL, command=self.filename_rules_tree.yview)
        self.filename_rules_tree.configure(yscrollcommand=rules_scroll.set)
        
        # Pack các widgets
        rules_scroll.pack(side=tk.RIGHT, fill=tk.Y)
        self.filename_rules_tree.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        
        # Frame cho phần chọn files/thư mục cần xử lý
        target_frame = ttk.LabelFrame(parent, text="Bước 2: Chọn Files hoặc Thư mục cần đổi tên", padding="10")
        target_frame.pack(fill=tk.X, pady=5)
        
        # Nút chọn files
        select_target_files_button = ttk.Button(target_frame, text="Chọn Files", command=self.select_target_files_for_rename)
        select_target_files_button.grid(row=0, column=0, padx=5, pady=5, sticky="w")
        
        # Nút chọn thư mục
        select_target_dir_button = ttk.Button(target_frame, text="Chọn Thư mục", command=self.select_target_directory_for_rename)
        select_target_dir_button.grid(row=0, column=1, padx=5, pady=5, sticky="w")
        
        # Hiển thị đường dẫn đã chọn
        self.filename_target_path_var = tk.StringVar()
        target_path_label = ttk.Label(target_frame, textvariable=self.filename_target_path_var, wraplength=800)
        target_path_label.grid(row=1, column=0, columnspan=2, padx=5, pady=5, sticky="w")
        
        # Tùy chọn include subdirectories
        self.filename_table_include_subdirs_var = tk.BooleanVar(value=True)
        include_subdirs_check = ttk.Checkbutton(
            target_frame, 
            text="Bao gồm các thư mục con",
            variable=self.filename_table_include_subdirs_var
        )
        include_subdirs_check.grid(row=2, column=0, columnspan=2, padx=5, pady=5, sticky="w")
        
        # Lựa chọn loại file
        ttk.Label(target_frame, text="Xử lý file:").grid(row=3, column=0, padx=5, pady=5, sticky="w")
        self.filename_table_type_frame = ttk.Frame(target_frame)
        self.filename_table_type_frame.grid(row=3, column=1, padx=5, pady=5, sticky="w")
        
        self.filename_table_process_docx_var = tk.BooleanVar(value=True)
        self.filename_table_process_doc_var = tk.BooleanVar(value=True)
        self.filename_table_process_all_var = tk.BooleanVar(value=False)
        
        ttk.Checkbutton(self.filename_table_type_frame, text=".docx", variable=self.filename_table_process_docx_var).pack(side=tk.LEFT, padx=5)
        ttk.Checkbutton(self.filename_table_type_frame, text=".doc", variable=self.filename_table_process_doc_var).pack(side=tk.LEFT, padx=5)
        ttk.Checkbutton(self.filename_table_type_frame, text="Tất cả các file", variable=self.filename_table_process_all_var, 
                        command=self.toggle_table_all_files).pack(side=tk.LEFT, padx=5)
        
        # Tùy chọn xem trước thay đổi
        preview_frame = ttk.LabelFrame(parent, text="Xem trước kết quả", padding="10")
        preview_frame.pack(fill=tk.BOTH, expand=True, pady=5)
        
        # Tạo Treeview để hiển thị danh sách tên file
        columns = ('old_name', 'new_name', 'path')
        self.filename_table_preview_tree = ttk.Treeview(preview_frame, columns=columns, show='headings', height=6)
        
        # Định nghĩa headings
        self.filename_table_preview_tree.heading('old_name', text='Tên cũ')
        self.filename_table_preview_tree.heading('new_name', text='Tên mới')
        self.filename_table_preview_tree.heading('path', text='Đường dẫn')
        
        # Định nghĩa columns
        self.filename_table_preview_tree.column('old_name', width=300)
        self.filename_table_preview_tree.column('new_name', width=300)
        self.filename_table_preview_tree.column('path', width=300)
        
        # Thêm thanh cuộn
        preview_scroll = ttk.Scrollbar(preview_frame, orient=tk.VERTICAL, command=self.filename_table_preview_tree.yview)
        self.filename_table_preview_tree.configure(yscrollcommand=preview_scroll.set)
        
        # Pack các widgets
        preview_scroll.pack(side=tk.RIGHT, fill=tk.Y)
        self.filename_table_preview_tree.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        
        # Nút xem trước và thực hiện
        button_frame = ttk.Frame(parent)
        button_frame.pack(fill=tk.X, pady=10)
        
        self.filename_table_preview_button = ttk.Button(
            button_frame, 
            text="Xem trước thay đổi", 
            command=self.preview_filename_table_changes
        )
        self.filename_table_preview_button.pack(side=tk.LEFT, padx=5)
        
        self.filename_table_execute_button = ttk.Button(
            button_frame, 
            text="Thực hiện đổi tên", 
            command=self.execute_filename_table_replace
        )
        self.filename_table_execute_button.pack(side=tk.RIGHT, padx=5)
        
        # Thanh tiến trình
        progress_frame = ttk.Frame(parent)
        progress_frame.pack(fill=tk.X, pady=5)
        
        ttk.Label(progress_frame, text="Tiến trình:").pack(side=tk.LEFT, padx=5)
        self.filename_table_progress = ttk.Progressbar(progress_frame, orient=tk.HORIZONTAL, length=800, mode='determinate')
        self.filename_table_progress.pack(side=tk.LEFT, padx=5, fill=tk.X, expand=True)
        
        # Trạng thái
        status_frame = ttk.Frame(parent)
        status_frame.pack(fill=tk.X, pady=5)
        
        self.filename_table_status_var = tk.StringVar(value="Sẵn sàng")
        status_label = ttk.Label(status_frame, textvariable=self.filename_table_status_var, wraplength=800)
        status_label.pack(side=tk.LEFT, padx=5, fill=tk.X, expand=True)
    def select_files(self):
        """Mở hộp thoại chọn nhiều file .doc và .docx"""
        file_types = []
        if self.process_docx_var.get():
            file_types.append("*.docx")
        if self.process_doc_var.get():
            file_types.append("*.doc")
            
        if not file_types:
            messagebox.showwarning("Cảnh báo", "Hãy chọn ít nhất một loại file cần xử lý (.doc hoặc .docx)")
            return
            
        files = filedialog.askopenfilenames(
            title="Chọn tệp Word",
            filetypes=[("Word files", " ".join(file_types))],
            multiple=True
        )
        if files:
            self.selected_files = list(files)
            self.selected_directory = ""
            self.path_var.set(f"Đã chọn {len(self.selected_files)} tệp")
            self.add_log(f"Đã chọn {len(self.selected_files)} tệp")
    
    def select_directory(self):
        """Mở hộp thoại chọn thư mục"""
        directory = filedialog.askdirectory(title="Chọn thư mục chứa tệp Word")
        if directory:
            self.selected_directory = directory
            self.selected_files = []
            self.path_var.set(f"Thư mục: {directory}")
            self.add_log(f"Đã chọn thư mục: {directory}")
    
    def get_file_filters_for_rename(self, docx_var, doc_var, all_var):
        """Trả về bộ lọc file cho việc chọn file đổi tên"""
        if all_var.get():
            return [("All files", "*.*")]
        
        file_types = []
        filters = []
        
        if docx_var.get():
            filters.append("*.docx")
            file_types.append("Word DOCX")
        
        if doc_var.get():
            filters.append("*.doc")
            file_types.append("Word DOC")
            
        if not filters:
            return []
            
        return [(f"{' & '.join(file_types)} files", " ".join(filters))]
    
    def toggle_all_files(self):
        """Xử lý khi người dùng chọn 'Tất cả các file'"""
        if self.filename_process_all_var.get():
            self.filename_process_docx_var.set(False)
            self.filename_process_doc_var.set(False)
        else:
            self.filename_process_docx_var.set(True)
            self.filename_process_doc_var.set(True)
    
    def toggle_table_all_files(self):
        """Xử lý khi người dùng chọn 'Tất cả các file' trong tab thay thế từ bảng"""
        if self.filename_table_process_all_var.get():
            self.filename_table_process_docx_var.set(False)
            self.filename_table_process_doc_var.set(False)
        else:
            self.filename_table_process_docx_var.set(True)
            self.filename_table_process_doc_var.set(True)
    
    def execute_search_replace(self):
        """Thực hiện tìm kiếm và thay thế đơn giản"""
        self.search_text = self.search_entry.get()
        self.replace_text = self.replace_entry.get()
        
        if not self.search_text:
            messagebox.showerror("Lỗi", "Vui lòng nhập chuỗi cần tìm!")
            return
        
        if not (self.selected_files or self.selected_directory):
            messagebox.showerror("Lỗi", "Vui lòng chọn files hoặc thư mục trước!")
            return
            
        if not (self.process_docx_var.get() or self.process_doc_var.get()):
            messagebox.showwarning("Cảnh báo", "Hãy chọn ít nhất một loại file cần xử lý (.doc hoặc .docx)")
            return
        
        # Xóa nội dung log cũ
        self.log_text.delete(1.0, tk.END)
        self.add_log(f"Bắt đầu tìm kiếm: '{self.search_text}' và thay thế bằng '{self.replace_text}'")
        
        # Reset danh sách files lỗi và số lượng
        self.failed_files = []
        self.success_count = 0
        self.skip_count = 0
        
        # Disable nút thực hiện để tránh click nhiều lần
        self.execute_button.config(state=tk.DISABLED)
        self.status_var.set("Đang xử lý...")
        
        # Chạy xử lý trong thread riêng để không làm đơ giao diện
        thread = threading.Thread(target=self.process_files_thread, args=([self.search_text], [self.replace_text]))
        thread.start()
    
    def execute_table_search_replace(self):
        """Thực hiện tìm kiếm và thay thế từ bảng quy tắc"""
        if not self.replacement_pairs:
            messagebox.showerror("Lỗi", "Vui lòng tải quy tắc thay thế từ file trước!")
            return
        
        if not (self.selected_files or self.selected_directory):
            messagebox.showerror("Lỗi", "Vui lòng chọn files hoặc thư mục cần áp dụng quy tắc!")
            return
            
        if not (self.target_process_docx_var.get() or self.target_process_doc_var.get()):
            messagebox.showwarning("Cảnh báo", "Hãy chọn ít nhất một loại file cần xử lý (.doc hoặc .docx)")
            return
        
        # Trích xuất danh sách tìm kiếm và thay thế từ replacement_pairs
        search_texts = [pair[0] for pair in self.replacement_pairs]
        replace_texts = [pair[1] for pair in self.replacement_pairs]
        
        # Xóa nội dung log cũ
        self.log_text.delete(1.0, tk.END)
        self.add_log(f"Bắt đầu áp dụng {len(self.replacement_pairs)} quy tắc thay thế")
        
        # Reset danh sách files lỗi và số lượng
        self.failed_files = []
        self.success_count = 0
        self.skip_count = 0
        
        # Disable nút thực hiện để tránh click nhiều lần
        self.target_execute_button.config(state=tk.DISABLED)
        self.target_status_var.set("Đang xử lý...")
        
        # Chạy xử lý trong thread riêng để không làm đơ giao diện
        thread = threading.Thread(target=self.process_files_thread, args=(search_texts, replace_texts, True))
        thread.start()
    
    def process_files_thread(self, search_texts, replace_texts, use_target_widgets=False):
        """Xử lý files trong thread riêng"""
        # Chọn widgets phù hợp dựa vào tab đang sử dụng
        if use_target_widgets:
            progress = self.target_progress
            status_var = self.target_status_var
            execute_button = self.target_execute_button
            process_docx_var = self.target_process_docx_var
            process_doc_var = self.target_process_doc_var
        else:
            progress = self.progress
            status_var = self.status_var
            execute_button = self.execute_button
            process_docx_var = self.process_docx_var
            process_doc_var = self.process_doc_var
        
        # Khởi tạo COM cho thread mới
        pythoncom.CoInitialize()
        
        try:
            if self.selected_files:
                # Xử lý các file đã chọn
                total = len(self.selected_files)
                for i, file_path in enumerate(self.selected_files):
                    self.update_progress((i / total) * 100, progress)
                    self.update_status(f"Đang xử lý: {os.path.basename(file_path)} ({i+1}/{total})", status_var)
                    self.process_single_file(file_path, search_texts, replace_texts, process_docx_var, process_doc_var)
            else:
                # Xử lý các file trong thư mục
                all_files = []
                for root, _, files in os.walk(self.selected_directory):
                    for file in files:
                        # Chỉ xử lý các loại file được chọn
                        if (file.lower().endswith('.docx') and process_docx_var.get()) or \
                           (file.lower().endswith('.doc') and process_doc_var.get()):
                            all_files.append(os.path.join(root, file))
                
                if not all_files:
                    self.add_log("Không tìm thấy file Word nào trong thư mục")
                    messagebox.showinfo("Thông báo", "Không tìm thấy file Word nào trong thư mục!")
                    return
                
                total = len(all_files)
                for i, file_path in enumerate(all_files):
                    self.update_progress((i / total) * 100, progress)
                    self.update_status(f"Đang xử lý: {os.path.basename(file_path)} ({i+1}/{total})", status_var)
                    self.process_single_file(file_path, search_texts, replace_texts, process_docx_var, process_doc_var)
            
            self.update_progress(100, progress)
            
            summary = f"\n===== Tổng kết =====\n"
            summary += f"- Đã xử lý thành công: {self.success_count} files\n"
            summary += f"- Không tìm thấy chuỗi cần thay: {self.skip_count} files\n"
            summary += f"- Gặp lỗi: {len(self.failed_files)} files\n"
            
            self.add_log(summary)
            
            if self.failed_files:
                failed_msg = f"Hoàn thành với {len(self.failed_files)} lỗi"
                self.update_status(failed_msg, status_var)
                self.add_log("Các file lỗi:")
                for file in self.failed_files:
                    self.add_log(f"- {file}")
                messagebox.showwarning("Cảnh báo", f"Đã hoàn tất với {len(self.failed_files)} lỗi. Xem log để biết chi tiết.")
            else:
                self.update_status("Hoàn thành!", status_var)
                messagebox.showinfo("Thông báo", "Quá trình tìm kiếm và thay thế đã hoàn tất!")
        
        except Exception as e:
            error_msg = f"Lỗi: {str(e)}"
            self.update_status(error_msg, status_var)
            self.add_log(f"\n{error_msg}")
            messagebox.showerror("Lỗi", f"Xảy ra lỗi: {str(e)}")
        
        finally:
            # Giải phóng COM
            pythoncom.CoUninitialize()
            # Enable lại nút thực hiện
            self.root.after(0, lambda: execute_button.config(state=tk.NORMAL))
    
    def execute_search_replace(self):
        """Thực hiện tìm kiếm và thay thế đơn giản"""
        self.search_text = self.search_entry.get()
        self.replace_text = self.replace_entry.get()
        
        if not self.search_text:
            messagebox.showerror("Lỗi", "Vui lòng nhập chuỗi cần tìm!")
            return
        
        if not (self.selected_files or self.selected_directory):
            messagebox.showerror("Lỗi", "Vui lòng chọn files hoặc thư mục trước!")
            return
            
        if not (self.process_docx_var.get() or self.process_doc_var.get()):
            messagebox.showwarning("Cảnh báo", "Hãy chọn ít nhất một loại file cần xử lý (.doc hoặc .docx)")
            return
        
        # Xóa nội dung log cũ
        self.log_text.delete(1.0, tk.END)
        self.add_log(f"Bắt đầu tìm kiếm: '{self.search_text}' và thay thế bằng '{self.replace_text}'")
        
        # Reset danh sách files lỗi và số lượng
        self.failed_files = []
        self.success_count = 0
        self.skip_count = 0
        
        # Disable nút thực hiện để tránh click nhiều lần
        self.execute_button.config(state=tk.DISABLED)
        self.status_var.set("Đang xử lý...")
        
        # Chạy xử lý trong thread riêng để không làm đơ giao diện
        thread = threading.Thread(target=self.process_files_thread, args=([self.search_text], [self.replace_text]))
        thread.start()
    
    def execute_table_search_replace(self):
        """Thực hiện tìm kiếm và thay thế từ bảng quy tắc"""
        if not self.replacement_pairs:
            messagebox.showerror("Lỗi", "Vui lòng tải quy tắc thay thế từ file trước!")
            return
        
        if not (self.selected_files or self.selected_directory):
            messagebox.showerror("Lỗi", "Vui lòng chọn files hoặc thư mục cần áp dụng quy tắc!")
            return
            
        if not (self.target_process_docx_var.get() or self.target_process_doc_var.get()):
            messagebox.showwarning("Cảnh báo", "Hãy chọn ít nhất một loại file cần xử lý (.doc hoặc .docx)")
            return
        
        # Trích xuất danh sách tìm kiếm và thay thế từ replacement_pairs
        search_texts = [pair[0] for pair in self.replacement_pairs]
        replace_texts = [pair[1] for pair in self.replacement_pairs]
        
        # Xóa nội dung log cũ
        self.log_text.delete(1.0, tk.END)
        self.add_log(f"Bắt đầu áp dụng {len(self.replacement_pairs)} quy tắc thay thế")
        
        # Reset danh sách files lỗi và số lượng
        self.failed_files = []
        self.success_count = 0
        self.skip_count = 0
        
        # Disable nút thực hiện để tránh click nhiều lần
        self.target_execute_button.config(state=tk.DISABLED)
        self.target_status_var.set("Đang xử lý...")
        
        # Chạy xử lý trong thread riêng để không làm đơ giao diện
        thread = threading.Thread(target=self.process_files_thread, args=(search_texts, replace_texts, True))
        thread.start()
    
    def process_files_thread(self, search_texts, replace_texts, use_target_widgets=False):
        """Xử lý files trong thread riêng"""
        # Chọn widgets phù hợp dựa vào tab đang sử dụng
        if use_target_widgets:
            progress = self.target_progress
            status_var = self.target_status_var
            execute_button = self.target_execute_button
            process_docx_var = self.target_process_docx_var
            process_doc_var = self.target_process_doc_var
        else:
            progress = self.progress
            status_var = self.status_var
            execute_button = self.execute_button
            process_docx_var = self.process_docx_var
            process_doc_var = self.process_doc_var
        
        # Khởi tạo COM cho thread mới
        pythoncom.CoInitialize()
        
        try:
            if self.selected_files:
                # Xử lý các file đã chọn
                total = len(self.selected_files)
                for i, file_path in enumerate(self.selected_files):
                    self.update_progress((i / total) * 100, progress)
                    self.update_status(f"Đang xử lý: {os.path.basename(file_path)} ({i+1}/{total})", status_var)
                    self.process_single_file(file_path, search_texts, replace_texts, process_docx_var, process_doc_var)
            else:
                # Xử lý các file trong thư mục
                all_files = []
                for root, _, files in os.walk(self.selected_directory):
                    for file in files:
                        # Chỉ xử lý các loại file được chọn
                        if (file.lower().endswith('.docx') and process_docx_var.get()) or \
                           (file.lower().endswith('.doc') and process_doc_var.get()):
                            all_files.append(os.path.join(root, file))
                
                if not all_files:
                    self.add_log("Không tìm thấy file Word nào trong thư mục")
                    messagebox.showinfo("Thông báo", "Không tìm thấy file Word nào trong thư mục!")
                    return
                
                total = len(all_files)
                for i, file_path in enumerate(all_files):
                    self.update_progress((i / total) * 100, progress)
                    self.update_status(f"Đang xử lý: {os.path.basename(file_path)} ({i+1}/{total})", status_var)
                    self.process_single_file(file_path, search_texts, replace_texts, process_docx_var, process_doc_var)
            
            self.update_progress(100, progress)
            
            summary = f"\n===== Tổng kết =====\n"
            summary += f"- Đã xử lý thành công: {self.success_count} files\n"
            summary += f"- Không tìm thấy chuỗi cần thay: {self.skip_count} files\n"
            summary += f"- Gặp lỗi: {len(self.failed_files)} files\n"
            
            self.add_log(summary)
            
            if self.failed_files:
                failed_msg = f"Hoàn thành với {len(self.failed_files)} lỗi"
                self.update_status(failed_msg, status_var)
                self.add_log("Các file lỗi:")
                for file in self.failed_files:
                    self.add_log(f"- {file}")
                messagebox.showwarning("Cảnh báo", f"Đã hoàn tất với {len(self.failed_files)} lỗi. Xem log để biết chi tiết.")
            else:
                self.update_status("Hoàn thành!", status_var)
                messagebox.showinfo("Thông báo", "Quá trình tìm kiếm và thay thế đã hoàn tất!")
        
        except Exception as e:
            error_msg = f"Lỗi: {str(e)}"
            self.update_status(error_msg, status_var)
            self.add_log(f"\n{error_msg}")
            messagebox.showerror("Lỗi", f"Xảy ra lỗi: {str(e)}")
        
        finally:
            # Giải phóng COM
            pythoncom.CoUninitialize()
            # Enable lại nút thực hiện
            self.root.after(0, lambda: execute_button.config(state=tk.NORMAL))
    
    def select_files_for_rename(self):
        """Mở hộp thoại chọn nhiều file để đổi tên"""
        file_types = self.get_file_filters_for_rename(self.filename_process_docx_var, 
                                                     self.filename_process_doc_var, 
                                                     self.filename_process_all_var)
            
        if not file_types:
            messagebox.showwarning("Cảnh báo", "Hãy chọn ít nhất một loại file cần xử lý")
            return
            
        files = filedialog.askopenfilenames(
            title="Chọn các tệp để đổi tên",
            filetypes=file_types,
            multiple=True
        )
        if files:
            self.selected_files = list(files)
            self.selected_directory = ""
            self.filename_path_var.set(f"Đã chọn {len(self.selected_files)} tệp")
            self.add_log(f"Đã chọn {len(self.selected_files)} tệp để đổi tên")
    
    def select_directory_for_rename(self):
        """Mở hộp thoại chọn thư mục cho việc đổi tên file"""
        directory = filedialog.askdirectory(title="Chọn thư mục chứa tệp cần đổi tên")
        if directory:
            self.selected_directory = directory
            self.selected_files = []
            self.filename_path_var.set(f"Thư mục: {directory}")
            self.add_log(f"Đã chọn thư mục: {directory} để tìm files đổi tên")
    
    def select_target_files_for_rename(self):
        """Mở hộp thoại chọn nhiều file để áp dụng quy tắc đổi tên"""
        file_types = self.get_file_filters_for_rename(self.filename_table_process_docx_var, 
                                                     self.filename_table_process_doc_var, 
                                                     self.filename_table_process_all_var)
            
        if not file_types:
            messagebox.showwarning("Cảnh báo", "Hãy chọn ít nhất một loại file cần xử lý")
            return
            
        files = filedialog.askopenfilenames(
            title="Chọn các tệp để đổi tên",
            filetypes=file_types,
            multiple=True
        )
        if files:
            self.selected_files = list(files)
            self.selected_directory = ""
            self.filename_target_path_var.set(f"Đã chọn {len(self.selected_files)} tệp")
            self.add_log(f"Đã chọn {len(self.selected_files)} tệp để áp dụng quy tắc đổi tên")
    
    def select_target_directory_for_rename(self):
        """Mở hộp thoại chọn thư mục để áp dụng quy tắc đổi tên"""
        directory = filedialog.askdirectory(title="Chọn thư mục chứa tệp cần đổi tên")
        if directory:
            self.selected_directory = directory
            self.selected_files = []
            self.filename_target_path_var.set(f"Thư mục: {directory}")
            self.add_log(f"Đã chọn thư mục: {directory} để áp dụng quy tắc đổi tên")
            
    def get_files_for_rename(self, process_docx_var, process_doc_var, process_all_var, include_subdirs_var=None):
        """Lấy danh sách các file cần đổi tên"""
        files_to_rename = []
        
        if self.selected_files:
            # Nếu người dùng đã chọn các file riêng lẻ
            for file_path in self.selected_files:
                if self.is_file_type_match(file_path, process_docx_var, process_doc_var, process_all_var):
                    files_to_rename.append(file_path)
        else:
            # Nếu người dùng đã chọn thư mục
            if include_subdirs_var and include_subdirs_var.get():
                # Bao gồm thư mục con
                for root, _, files in os.walk(self.selected_directory):
                    for file in files:
                        file_path = os.path.join(root, file)
                        if self.is_file_type_match(file_path, process_docx_var, process_doc_var, process_all_var):
                            files_to_rename.append(file_path)
            else:
                # Chỉ thư mục hiện tại
                for file in os.listdir(self.selected_directory):
                    file_path = os.path.join(self.selected_directory, file)
                    if os.path.isfile(file_path) and self.is_file_type_match(file_path, process_docx_var, process_doc_var, process_all_var):
                        files_to_rename.append(file_path)
        
        return files_to_rename
    
    def is_file_type_match(self, file_path, process_docx_var, process_doc_var, process_all_var):
        """Kiểm tra xem file có khớp với loại file được chọn hay không"""
        if process_all_var.get():
            return True
            
        lower_path = file_path.lower()
        if lower_path.endswith('.docx') and process_docx_var.get():
            return True
        elif lower_path.endswith('.doc') and process_doc_var.get():
            return True
        
        return False
            
   
    def preview_filename_changes(self):
        """Xem trước thay đổi tên file"""
        search_text = self.filename_search_entry.get()
        replace_text = self.filename_replace_entry.get()
        
        if not search_text:
            messagebox.showerror("Lỗi", "Vui lòng nhập chuỗi cần tìm trong tên file!")
            return
        
        if not (self.selected_files or self.selected_directory):
            messagebox.showerror("Lỗi", "Vui lòng chọn files hoặc thư mục trước!")
            return
            
        if not (self.filename_process_docx_var.get() or self.filename_process_doc_var.get() or self.filename_process_all_var.get()):
            messagebox.showwarning("Cảnh báo", "Hãy chọn ít nhất một loại file cần xử lý")
            return
        
        # Xóa dữ liệu cũ trong bảng xem trước
        for item in self.filename_preview_tree.get_children():
            self.filename_preview_tree.delete(item)
        
        # Cập nhật trạng thái
        self.filename_status_var.set("Đang tìm kiếm files...")
        
        # Lấy danh sách các file cần đổi tên
        files_to_rename = self.get_files_for_rename(
            self.filename_process_docx_var, 
            self.filename_process_doc_var, 
            self.filename_process_all_var,
            self.filename_include_subdirs_var
        )
        
        # Hiển thị kết quả xem trước
        for file_path in files_to_rename:
            old_name = os.path.basename(file_path)
            dir_name = os.path.dirname(file_path)
            
            # Tách phần tên và phần mở rộng
            name_parts = os.path.splitext(old_name)
            file_name = name_parts[0]
            file_ext = name_parts[1] if len(name_parts) > 1 else ""
            
            # Thay thế chỉ trong phần tên file, không đụng đến phần mở rộng
            new_name = file_name.replace(search_text, replace_text) + file_ext
            
            # Chỉ hiển thị các file thực sự có thay đổi
            if old_name != new_name:
                self.filename_preview_tree.insert('', 'end', values=(old_name, new_name, dir_name))
                
        # Cập nhật trạng thái
        item_count = len(self.filename_preview_tree.get_children())
        if item_count > 0:
            self.filename_status_var.set(f"Tìm thấy {item_count} file cần đổi tên")
        else:
            self.filename_status_var.set("Không tìm thấy file nào khớp với điều kiện tìm kiếm")
    
    def execute_filename_replace(self):
        """Thực hiện đổi tên file dựa trên kết quả xem trước"""
        # Kiểm tra xem có files được xem trước không
        if not self.filename_preview_tree.get_children():
            messagebox.showinfo("Thông báo", "Không có file nào cần đổi tên hoặc bạn chưa xem trước thay đổi!")
            return
        
        # Xác nhận từ người dùng
        confirm = messagebox.askyesno(
            "Xác nhận", 
            "Bạn có chắc chắn muốn đổi tên các file đã chọn?\n\nLưu ý: Quá trình này không thể hoàn tác!"
        )
        if not confirm:
            return
        
        # Thực hiện đổi tên
        success_count = 0
        failed_files = []
        
        # Disable nút để tránh nhấn nhiều lần
        self.filename_execute_button.config(state=tk.DISABLED)
        self.filename_preview_button.config(state=tk.DISABLED)
        
        # Cập nhật trạng thái
        self.filename_status_var.set("Đang đổi tên files...")
        
        # Lặp qua tất cả các mục trong danh sách xem trước
        items = self.filename_preview_tree.get_children()
        total = len(items)
        
        for i, item in enumerate(items):
            values = self.filename_preview_tree.item(item, 'values')
            old_name = values[0]
            new_name = values[1]
            dir_path = values[2]
            
            old_path = os.path.join(dir_path, old_name)
            new_path = os.path.join(dir_path, new_name)
            
            # Cập nhật tiến trình
            self.update_progress((i / total) * 100, self.filename_progress)
            self.update_status(f"Đang đổi tên: {old_name} -> {new_name}", self.filename_status_var)
            
            try:
                # Kiểm tra xem file mới đã tồn tại chưa
                if os.path.exists(new_path):
                    self.add_log(f"✗ Lỗi: File '{new_name}' đã tồn tại trong thư mục '{dir_path}'")
                    failed_files.append(old_path)
                    continue
                
                # Thực hiện đổi tên
                os.rename(old_path, new_path)
                success_count += 1
                self.add_log(f"✓ Đã đổi tên: {old_name} -> {new_name}")
            
            except Exception as e:
                self.add_log(f"✗ Lỗi khi đổi tên '{old_name}': {str(e)}")
                failed_files.append(old_path)
        
        # Cập nhật tiến trình hoàn thành
        self.update_progress(100, self.filename_progress)
        
        # Hiển thị thông báo kết quả
        summary = f"\n===== Tổng kết đổi tên file =====\n"
        summary += f"- Đã đổi tên thành công: {success_count} files\n"
        summary += f"- Gặp lỗi: {len(failed_files)} files\n"
        
        self.add_log(summary)
        
        if failed_files:
            self.filename_status_var.set(f"Hoàn thành với {len(failed_files)} lỗi")
            messagebox.showwarning("Cảnh báo", f"Đã hoàn tất với {len(failed_files)} lỗi. Xem log để biết chi tiết.")
        else:
            self.filename_status_var.set(f"Đã đổi tên thành công {success_count} files")
            messagebox.showinfo("Thông báo", f"Đã đổi tên thành công {success_count} files!")
        
        # Xóa danh sách xem trước sau khi hoàn thành
        for item in self.filename_preview_tree.get_children():
            self.filename_preview_tree.delete(item)
        
        # Enable lại các nút
        self.filename_execute_button.config(state=tk.NORMAL)
        self.filename_preview_button.config(state=tk.NORMAL)
    
    def preview_filename_table_changes(self):
        """Xem trước thay đổi tên file từ bảng quy tắc"""
        if not self.filename_replacement_pairs:
            messagebox.showerror("Lỗi", "Vui lòng tải quy tắc thay thế tên file từ file trước!")
            return
        
        if not (self.selected_files or self.selected_directory):
            messagebox.showerror("Lỗi", "Vui lòng chọn files hoặc thư mục cần áp dụng quy tắc!")
            return
            
        if not (self.filename_table_process_docx_var.get() or self.filename_table_process_doc_var.get() or self.filename_table_process_all_var.get()):
            messagebox.showwarning("Cảnh báo", "Hãy chọn ít nhất một loại file cần xử lý")
            return
        
        # Xóa dữ liệu cũ trong bảng xem trước
        for item in self.filename_table_preview_tree.get_children():
            self.filename_table_preview_tree.delete(item)
        
        # Cập nhật trạng thái
        self.filename_table_status_var.set("Đang tìm kiếm files...")
        
        # Lấy danh sách các file cần đổi tên
        files_to_rename = self.get_files_for_rename(
            self.filename_table_process_docx_var, 
            self.filename_table_process_doc_var, 
            self.filename_table_process_all_var,
            self.filename_table_include_subdirs_var
        )
        
        # Hiển thị kết quả xem trước
        for file_path in files_to_rename:
            old_name = os.path.basename(file_path)
            dir_name = os.path.dirname(file_path)
            
            # Tách phần tên và phần mở rộng
            name_parts = os.path.splitext(old_name)
            file_name = name_parts[0]
            file_ext = name_parts[1] if len(name_parts) > 1 else ""
            
            # Áp dụng tất cả các quy tắc
            new_file_name = file_name
            changed = False
            
            for search_text, replace_text in self.filename_replacement_pairs:
                if search_text in new_file_name:
                    new_file_name = new_file_name.replace(search_text, replace_text)
                    changed = True
                    
            # Tạo tên file mới đầy đủ
            new_name = new_file_name + file_ext
            
            # Chỉ hiển thị các file thực sự có thay đổi
            if changed and old_name != new_name:
                self.filename_table_preview_tree.insert('', 'end', values=(old_name, new_name, dir_name))
                
        # Cập nhật trạng thái
        item_count = len(self.filename_table_preview_tree.get_children())
        if item_count > 0:
            self.filename_table_status_var.set(f"Tìm thấy {item_count} file cần đổi tên")
        else:
            self.filename_table_status_var.set("Không tìm thấy file nào khớp với quy tắc thay thế")
    
    def select_rules_file(self):
        """Mở hộp thoại chọn file chứa bảng quy tắc thay thế"""
        file_path = filedialog.askopenfilename(
            title="Chọn tệp chứa bảng quy tắc",
            filetypes=[("Word/Excel files", "*.docx *.xlsx *.xls")]
        )
        if file_path:
            self.rules_file_path = file_path
            self.rules_path_var.set(f"Tệp quy tắc: {file_path}")
            self.load_replacement_rules(file_path)

    def load_replacement_rules(self, file_path):
        """Tải quy tắc thay thế từ file Word hoặc Excel"""
        self.replacement_pairs = []
        
        try:
            # Xóa dữ liệu cũ trong bảng quy tắc
            for item in self.rules_tree.get_children():
                self.rules_tree.delete(item)
            
            if file_path.lower().endswith('.docx'):
                # Đọc từ file Word
                doc = Document(file_path)
                
                # Tìm bảng đầu tiên trong tài liệu
                if len(doc.tables) > 0:
                    table = doc.tables[0]
                    
                    # Tìm cột "Tìm kiếm" và "Thay thế"
                    header_row = table.rows[0]
                    search_col_idx = -1
                    replace_col_idx = -1
                    
                    for i, cell in enumerate(header_row.cells):
                        if "tìm kiếm" in cell.text.lower():
                            search_col_idx = i
                        elif "thay thế" in cell.text.lower():
                            replace_col_idx = i
                    
                    if search_col_idx >= 0 and replace_col_idx >= 0:
                        # Đọc các quy tắc từ bảng
                        for i, row in enumerate(table.rows[1:], 1):  # Skip header row
                            if len(row.cells) > max(search_col_idx, replace_col_idx):
                                search_text = row.cells[search_col_idx].text.strip()
                                replace_text = row.cells[replace_col_idx].text.strip()
                                
                                if search_text:  # Chỉ thêm nếu ô tìm kiếm không trống
                                    self.replacement_pairs.append((search_text, replace_text))
                                    self.rules_tree.insert('', 'end', values=(i, search_text, replace_text))
                    else:
                        raise ValueError("Không tìm thấy cột 'Tìm kiếm' hoặc 'Thay thế' trong bảng")
                else:
                    raise ValueError("Không tìm thấy bảng trong tài liệu Word")
                    
            elif file_path.lower().endswith(('.xlsx', '.xls')):
                # Đọc từ file Excel
                # Sử dụng pandas để đọc file Excel
                df = pd.read_excel(file_path)
                
                # Tìm cột "Tìm kiếm" và "Thay thế"
                search_col = None
                replace_col = None
                
                for col in df.columns:
                    if isinstance(col, str) and "tìm kiếm" in col.lower():
                        search_col = col
                    elif isinstance(col, str) and "thay thế" in col.lower():
                        replace_col = col
                
                if search_col and replace_col:
                    # Đọc các quy tắc từ DataFrame
                    for i, (_, row) in enumerate(df.iterrows(), 1):
                        search_text = str(row[search_col]).strip()
                        replace_text = str(row[replace_col]).strip()
                        
                        # Kiểm tra giá trị NaN
                        if search_text.lower() == "nan":
                            search_text = ""
                        if replace_text.lower() == "nan":
                            replace_text = ""
                            
                        if search_text:  # Chỉ thêm nếu ô tìm kiếm không trống
                            self.replacement_pairs.append((search_text, replace_text))
                            self.rules_tree.insert('', 'end', values=(i, search_text, replace_text))
                else:
                    raise ValueError("Không tìm thấy cột 'Tìm kiếm' hoặc 'Thay thế' trong bảng Excel")
            
            self.add_log(f"Đã tải {len(self.replacement_pairs)} quy tắc thay thế từ file {os.path.basename(file_path)}")
            messagebox.showinfo("Thông báo", f"Đã tải {len(self.replacement_pairs)} quy tắc thay thế từ file.")
            
        except Exception as e:
            messagebox.showerror("Lỗi", f"Lỗi khi tải quy tắc thay thế: {str(e)}")
            self.rules_path_var.set("Lỗi tải tệp quy tắc")
            self.add_log(f"Lỗi khi tải quy tắc: {str(e)}")

    def select_filename_rules_file(self):
        """Mở hộp thoại chọn file chứa bảng quy tắc thay thế tên file"""
        file_path = filedialog.askopenfilename(
            title="Chọn tệp chứa bảng quy tắc đổi tên file",
            filetypes=[("Word/Excel files", "*.docx *.xlsx *.xls")]
        )
        if file_path:
            self.filename_rules_file_path = file_path
            self.filename_rules_path_var.set(f"Tệp quy tắc: {file_path}")
            self.load_filename_replacement_rules(file_path)

    def load_filename_replacement_rules(self, file_path):
        """Tải quy tắc thay thế tên file từ file Word hoặc Excel"""
        self.filename_replacement_pairs = []
        
        try:
            # Xóa dữ liệu cũ trong bảng quy tắc
            for item in self.filename_rules_tree.get_children():
                self.filename_rules_tree.delete(item)
            
            if file_path.lower().endswith('.docx'):
                # Đọc từ file Word
                doc = Document(file_path)
                
                # Tìm bảng đầu tiên trong tài liệu
                if len(doc.tables) > 0:
                    table = doc.tables[0]
                    
                    # Tìm cột "Tìm kiếm" và "Thay thế"
                    header_row = table.rows[0]
                    search_col_idx = -1
                    replace_col_idx = -1
                    
                    for i, cell in enumerate(header_row.cells):
                        if "tìm kiếm" in cell.text.lower():
                            search_col_idx = i
                        elif "thay thế" in cell.text.lower():
                            replace_col_idx = i
                    
                    if search_col_idx >= 0 and replace_col_idx >= 0:
                        # Đọc các quy tắc từ bảng
                        for i, row in enumerate(table.rows[1:], 1):  # Skip header row
                            if len(row.cells) > max(search_col_idx, replace_col_idx):
                                search_text = row.cells[search_col_idx].text.strip()
                                replace_text = row.cells[replace_col_idx].text.strip()
                                
                                if search_text:  # Chỉ thêm nếu ô tìm kiếm không trống
                                    self.filename_replacement_pairs.append((search_text, replace_text))
                                    self.filename_rules_tree.insert('', 'end', values=(i, search_text, replace_text))
                    else:
                        raise ValueError("Không tìm thấy cột 'Tìm kiếm' hoặc 'Thay thế' trong bảng")
                else:
                    raise ValueError("Không tìm thấy bảng trong tài liệu Word")
                    
            elif file_path.lower().endswith(('.xlsx', '.xls')):
                # Đọc từ file Excel
                df = pd.read_excel(file_path)
                
                # Tìm cột "Tìm kiếm" và "Thay thế"
                search_col = None
                replace_col = None
                
                for col in df.columns:
                    if isinstance(col, str) and "tìm kiếm" in col.lower():
                        search_col = col
                    elif isinstance(col, str) and "thay thế" in col.lower():
                        replace_col = col
                
                if search_col and replace_col:
                    # Đọc các quy tắc từ DataFrame
                    for i, (_, row) in enumerate(df.iterrows(), 1):
                        search_text = str(row[search_col]).strip()
                        replace_text = str(row[replace_col]).strip()
                        
                        # Kiểm tra giá trị NaN
                        if search_text.lower() == "nan":
                            search_text = ""
                        if replace_text.lower() == "nan":
                            replace_text = ""
                            
                        if search_text:  # Chỉ thêm nếu ô tìm kiếm không trống
                            self.filename_replacement_pairs.append((search_text, replace_text))
                            self.filename_rules_tree.insert('', 'end', values=(i, search_text, replace_text))
                else:
                    raise ValueError("Không tìm thấy cột 'Tìm kiếm' hoặc 'Thay thế' trong bảng Excel")
            
            self.add_log(f"Đã tải {len(self.filename_replacement_pairs)} quy tắc thay thế tên file từ file {os.path.basename(file_path)}")
            messagebox.showinfo("Thông báo", f"Đã tải {len(self.filename_replacement_pairs)} quy tắc thay thế tên file.")
            
        except Exception as e:
            messagebox.showerror("Lỗi", f"Lỗi khi tải quy tắc thay thế: {str(e)}")
            self.filename_rules_path_var.set("Lỗi tải tệp quy tắc")
            self.add_log(f"Lỗi khi tải quy tắc tên file: {str(e)}")

    def select_target_files(self):
        """Mở hộp thoại chọn nhiều file .doc và .docx để áp dụng quy tắc"""
        file_types = []
        if self.target_process_docx_var.get():
            file_types.append("*.docx")
        if self.target_process_doc_var.get():
            file_types.append("*.doc")
            
        if not file_types:
            messagebox.showwarning("Cảnh báo", "Hãy chọn ít nhất một loại file cần xử lý (.doc hoặc .docx)")
            return
            
        files = filedialog.askopenfilenames(
            title="Chọn tệp Word để áp dụng quy tắc",
            filetypes=[("Word files", " ".join(file_types))],
            multiple=True
        )
        if files:
            self.selected_files = list(files)
            self.selected_directory = ""
            self.target_path_var.set(f"Đã chọn {len(self.selected_files)} tệp")
            self.add_log(f"Đã chọn {len(self.selected_files)} tệp để áp dụng quy tắc")

    def select_target_directory(self):
        """Mở hộp thoại chọn thư mục để áp dụng quy tắc"""
        directory = filedialog.askdirectory(title="Chọn thư mục chứa tệp Word để áp dụng quy tắc")
        if directory:
            self.selected_directory = directory
            self.selected_files = []
            self.target_path_var.set(f"Thư mục: {directory}")
            self.add_log(f"Đã chọn thư mục: {directory} để áp dụng quy tắc")
    
    def execute_filename_table_replace(self):
        """Thực hiện đổi tên file dựa trên kết quả xem trước từ bảng quy tắc"""
        # Kiểm tra xem có files được xem trước không
        if not self.filename_table_preview_tree.get_children():
            messagebox.showinfo("Thông báo", "Không có file nào cần đổi tên hoặc bạn chưa xem trước thay đổi!")
            return
        
        # Xác nhận từ người dùng
        confirm = messagebox.askyesno(
            "Xác nhận", 
            "Bạn có chắc chắn muốn đổi tên các file đã chọn?\n\nLưu ý: Quá trình này không thể hoàn tác!"
        )
        if not confirm:
            return
        
        # Thực hiện đổi tên
        success_count = 0
        failed_files = []
        
        # Disable nút để tránh nhấn nhiều lần
        self.filename_table_execute_button.config(state=tk.DISABLED)
        self.filename_table_preview_button.config(state=tk.DISABLED)
        
        # Cập nhật trạng thái
        self.filename_table_status_var.set("Đang đổi tên files...")
        
        # Lặp qua tất cả các mục trong danh sách xem trước
        items = self.filename_table_preview_tree.get_children()
        total = len(items)
        
        for i, item in enumerate(items):
            values = self.filename_table_preview_tree.item(item, 'values')
            old_name = values[0]
            new_name = values[1]
            dir_path = values[2]
            
            old_path = os.path.join(dir_path, old_name)
            new_path = os.path.join(dir_path, new_name)
            
            # Cập nhật tiến trình
            self.update_progress((i / total) * 100, self.filename_table_progress)
            self.update_status(f"Đang đổi tên: {old_name} -> {new_name}", self.filename_table_status_var)
            
            try:
                # Kiểm tra xem file mới đã tồn tại chưa
                if os.path.exists(new_path):
                    self.add_log(f"✗ Lỗi: File '{new_name}' đã tồn tại trong thư mục '{dir_path}'")
                    failed_files.append(old_path)
                    continue
                
                # Thực hiện đổi tên
                os.rename(old_path, new_path)
                success_count += 1
                self.add_log(f"✓ Đã đổi tên: {old_name} -> {new_name}")
            
            except Exception as e:
                self.add_log(f"✗ Lỗi khi đổi tên '{old_name}': {str(e)}")
                failed_files.append(old_path)
        
        # Cập nhật tiến trình hoàn thành
        self.update_progress(100, self.filename_table_progress)
        
        # Hiển thị thông báo kết quả
        summary = f"\n===== Tổng kết đổi tên file =====\n"
        summary += f"- Đã đổi tên thành công: {success_count} files\n"
        summary += f"- Gặp lỗi: {len(failed_files)} files\n"
        
        self.add_log(summary)
        
        if failed_files:
            self.filename_table_status_var.set(f"Hoàn thành với {len(failed_files)} lỗi")
            messagebox.showwarning("Cảnh báo", f"Đã hoàn tất với {len(failed_files)} lỗi. Xem log để biết chi tiết.")
        else:
            self.filename_table_status_var.set(f"Đã đổi tên thành công {success_count} files")
            messagebox.showinfo("Thông báo", f"Đã đổi tên thành công {success_count} files!")
        
        # Xóa danh sách xem trước sau khi hoàn thành
        for item in self.filename_table_preview_tree.get_children():
            self.filename_table_preview_tree.delete(item)
        
        # Enable lại các nút
        self.filename_table_execute_button.config(state=tk.NORMAL)
        self.filename_table_preview_button.config(state=tk.NORMAL)
        
    def update_progress(self, value, progress_bar):
        """Cập nhật thanh tiến trình"""
        self.root.after(0, lambda: progress_bar.config(value=value))
    
    def update_status(self, text, status_var):
        """Cập nhật trạng thái"""
        self.root.after(0, lambda: status_var.set(text))
    
    def add_log(self, message):
        """Thêm thông báo vào vùng log"""
        self.root.after(0, lambda: self._add_log_internal(message))
        
    def _add_log_internal(self, message):
        self.log_text.insert(tk.END, message + "\n")
        self.log_text.see(tk.END)  # Cuộn xuống dòng mới nhất

# Main code to run the application
if __name__ == "__main__":
    # Thêm thư viện cần thiết
    required_libraries = ['python-docx', 'pywin32', 'pandas', 'openpyxl']
    
    # Kiểm tra các thư viện và cài đặt nếu cần
    try:
        import importlib
        missing_libs = []
        
        for lib in required_libraries:
            try:
                importlib.import_module(lib.replace('-', ''))
            except ImportError:
                missing_libs.append(lib)
        
        if missing_libs:
            print(f"Đang cài đặt các thư viện cần thiết: {', '.join(missing_libs)}")
            import subprocess
            for lib in missing_libs:
                subprocess.check_call([sys.executable, "-m", "pip", "install", lib])
            print("Đã cài đặt xong các thư viện.")
    except Exception as e:
        print(f"Lỗi khi cài đặt thư viện: {str(e)}")
        print("Vui lòng cài đặt thủ công các thư viện sau: python-docx, pywin32, pandas, openpyxl")
    
    root = tk.Tk()
    app = SearchReplaceApp(root)
    root.mainloop()